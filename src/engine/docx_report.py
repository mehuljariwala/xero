from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


REPORT_NAMES = {
    'trial_balance_report': 'Trial Balance',
    'profit_and_loss': 'Profit & Loss',
    'aged_receivables_detail': 'Aged Receivables',
    'aged_payables_detail': 'Aged Payables',
    'account_transactions': 'Account Transactions',
    'receivable_invoice_detail': 'Receivable Invoice Detail',
    'payable_invoice_detail': 'Payable Invoice Detail',
    'vat_returns_export': 'VAT Returns',
    'vat_returns': 'VAT Returns',
}

FRIENDLY_FILTER_NAMES = {
    'financial_year_start_date': 'Start Date',
    'financial_year_end_date': 'End Date',
    'report_end_date': 'End Date',
    'selected_client': 'Client',
    'company_name': 'Company',
    'vat_return_period': 'VAT Period',
    'vat_return_start_date': 'VAT Start',
    'vat_return_end_date': 'VAT End',
}

REPORT_WORKFLOW_NAMES = {
    'Trial Balance', 'Profit & Loss', 'Aged Receivables', 'Aged Payables',
    'Account Transactions', 'Receivable Invoice Detail', 'Payable Invoice Detail',
    'VAT Returns',
}

COLOR_HEADER_BG = RGBColor(0x1E, 0x29, 0x3B)
COLOR_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_SUCCESS = RGBColor(0x05, 0x96, 0x69)
COLOR_FAILED = RGBColor(0xDC, 0x26, 0x26)
COLOR_SKIPPED = RGBColor(0xD9, 0x77, 0x06)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_DARK = RGBColor(0x0F, 0x17, 0x2A)
COLOR_ROW_ALT = RGBColor(0xF1, 0xF5, 0xF9)


def _set_cell_shading(cell, color_hex: str):
    shading = cell._tc.get_or_add_tcPr()
    shading_el = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_el)


def _format_duration(seconds: float) -> str:
    if seconds >= 60:
        return f"{int(seconds // 60)}m {int(seconds % 60)}s"
    return f"{seconds:.1f}s"


def _extract_client_data(events: list) -> dict:
    clients = {}
    current_entry = None

    for event in events:
        if event["type"] == "workflow_start":
            if current_entry and current_entry["display_name"] in REPORT_WORKFLOW_NAMES:
                client_key = current_entry["client"] or "__global__"
                clients.setdefault(client_key, []).append(current_entry)

            workflow_name = event.get("workflow", "")
            display_name = REPORT_NAMES.get(workflow_name, workflow_name)

            current_entry = {
                "display_name": display_name,
                "client": event.get("client", ""),
                "start_time": event.get("time", ""),
                "end_time": "",
                "duration": 0,
                "filters": {},
                "downloaded": False,
                "skipped": False,
                "status": "running",
            }

        elif event["type"] == "filter" and current_entry:
            filter_name = event.get("filter_name", "")
            filter_value = event.get("value", "")
            friendly = FRIENDLY_FILTER_NAMES.get(filter_name, filter_name.replace('_', ' ').title())
            current_entry["filters"][friendly] = filter_value

        elif event["type"] == "download" and current_entry:
            current_entry["downloaded"] = True

        elif event["type"] == "skip" and current_entry:
            current_entry["skipped"] = True

        elif event["type"] == "error" and current_entry:
            if event.get("fatal"):
                current_entry["status"] = "failed"

        elif event["type"] == "workflow_end" and current_entry:
            current_entry["end_time"] = event.get("time", "")
            current_entry["duration"] = event.get("duration_seconds", 0)
            end_status = event.get("status", "completed")
            if current_entry["status"] != "failed":
                current_entry["status"] = end_status

    if current_entry and current_entry["display_name"] in REPORT_WORKFLOW_NAMES:
        client_key = current_entry["client"] or "__global__"
        clients.setdefault(client_key, []).append(current_entry)

    return clients


def generate_docx_report(events: list, output_path: str = None) -> str:
    if not output_path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"downloads/execution_summary_{timestamp}.docx"

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    title = doc.add_heading('Execution Summary Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_DARK
        run.font.size = Pt(22)

    chain_start = next((e for e in events if e["type"] == "workflow_start"), {})
    chain_end = None
    for e in reversed(events):
        if e["type"] == "workflow_end":
            chain_end = e
            break

    run_date = datetime.now().strftime("%d %B %Y")
    start_time = chain_start.get("time", "N/A")
    end_time = chain_end.get("time", "N/A") if chain_end else "N/A"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated: {run_date}  |  Run: {start_time} - {end_time}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    clients = _extract_client_data(events)

    if not clients:
        p = doc.add_paragraph()
        run = p.add_run("No report workflows were executed in this run.")
        run.font.color.rgb = COLOR_MUTED
        run.font.italic = True
        doc.save(output_path)
        return output_path

    headers = ['Report', 'Filters Applied', 'Start', 'End', 'Duration', 'Status']
    col_widths = [Inches(1.5), Inches(2.6), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]

    for client_name, report_entries in clients.items():
        display_name = client_name if client_name != "__global__" else "All Clients"

        client_heading = doc.add_heading(display_name, level=2)
        for run in client_heading.runs:
            run.font.color.rgb = COLOR_DARK
            run.font.size = Pt(14)

        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for i, (header, width) in enumerate(zip(headers, col_widths)):
            cell = table.rows[0].cells[i]
            cell.width = width
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = COLOR_HEADER_TEXT
            _set_cell_shading(cell, '1E293B')

        for row_idx, entry in enumerate(report_entries):
            filter_parts = [f"{k}: {v}" for k, v in entry["filters"].items() if k not in ('Client', 'Company')]
            filter_str = ", ".join(filter_parts) if filter_parts else "-"

            duration_str = _format_duration(entry["duration"])

            if entry["skipped"]:
                status_text = "SKIPPED"
                status_color = COLOR_SKIPPED
            elif entry["status"] == "failed":
                status_text = "FAILED"
                status_color = COLOR_FAILED
            elif entry["downloaded"] or entry["status"] == "completed":
                status_text = "SUCCESS"
                status_color = COLOR_SUCCESS
            else:
                status_text = "FAILED"
                status_color = COLOR_FAILED

            row_data = [
                entry["display_name"],
                filter_str,
                entry["start_time"],
                entry["end_time"],
                duration_str,
                status_text,
            ]

            row = table.add_row()

            if row_idx % 2 == 1:
                for cell in row.cells:
                    _set_cell_shading(cell, 'F1F5F9')

            for col_idx, value in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.width = col_widths[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]

                if col_idx in (2, 3, 4, 5):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = p.add_run(value)
                run.font.size = Pt(8.5)

                if col_idx == 0:
                    run.bold = True
                elif col_idx == 5:
                    run.bold = True
                    run.font.color.rgb = status_color

        doc.add_paragraph()

    doc.save(output_path)
    return output_path
